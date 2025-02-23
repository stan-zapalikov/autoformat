import regex as re
from docx import Document

def iterate_lines(doc_path):
    doc = Document(doc_path)
    for paragraph in doc.paragraphs:
        for line in paragraph.text.splitlines():
            yield line

def remove_digits_regex(text):
    return re.sub(r"^\s*\d+\s*$", "", text)

def remove_unknown_chars(text):
    return re.sub(r"[^\w\s\p{Punct}]+", "", text)

def numbers_to_superscript(document):
    for paragraph in document.paragraphs:
        for run in paragraph.runs:
            text = run.text
            if re.search(r"\d", text):
                new_runs = []
                parts = re.split(r"(\d+)", text)
                for part in parts:
                    new_run = paragraph.add_run(part)
                    if re.match(r"\d+", part):
                        # Check for adjacent punctuation or letters
                        if (
                            (parts.index(part) > 0 and re.search(r"[\w\p{Punct}]$", parts[parts.index(part) - 1]))
                            or (parts.index(part) < len(parts) - 1 and re.search(r"^[\w\p{Punct}]", parts[parts.index(part) + 1]))
                        ):
                            # Adjacent character found, don't make superscript
                            pass
                        else:
                            new_run.font.superscript = True
                    new_runs.append(new_run)
                run.text = ""
                for new_run in new_runs:
                    paragraph._p.append(new_run._r)
    return document

def remove_broken_lines(original_doc_path, new_doc_path):
    new_doc = Document()
    paragraph_text = ""
    for line in iterate_lines(original_doc_path):
        print(paragraph_text)
        if not line.rstrip().endswith('.') :
            paragraph_text += line.strip("\t\n")
        elif line.rstrip().endswith('.'):
            paragraph_text += line.strip("\t\n")
            paragraph_text += "\n"
            new_doc.add_paragraph(remove_unknown_chars(remove_digits_regex(paragraph_text)))
            paragraph_text = "\t"
    new_doc = numbers_to_superscript(new_doc)
    new_doc.save(new_doc_path)


remove_broken_lines("./documents/original.docx", "./documents/new.docx")
